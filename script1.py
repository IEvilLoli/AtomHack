# pip install pypiwin32

import docx
import os
import datetime
import re
import docx_parser
import os
import glob
import subprocess
from win32com import client as wc




# def packageformatdef(s):
#     length = len(s)
#     integers = []
#     i = 0  # индекс текущего символа
#
#     while i < length:
#         s_int = ''  # строка для нового числа
#         while i < length and '0' <= s[i] <= '9':
#             s_int += s[i]
#             i += 1
#         i += 1
#         if s_int != '':
#             integers.append(int(s_int))
#
#     package_format_str = '{0} {1} {2} {3}'.format(integers[0], integers[1], integers[2], integers[3])
#
#     return package_format_str


def createxml():
    print(0)


if __name__ == '__main__':
    # filepath = "C:/hack/Check-list_5 9 3 10.docx"

    # собираем все файлы в папке
    all_files = []
    for root, dirs, files in os.walk("data"):
        for filename in files:
            all_files.append(filename)
            print(filename)

    file_path = ' '.join(all_files)
    # Ищем только файл ведомости
    file_path = re.search(r"[R,r][^.WP]{1,}WP \S{1,}\d\.doc\S*", file_path)
    print("find")
    print(file_path[0] if file_path else 'Not found')
    filepath = "data/" + file_path[0] if file_path else 'Not found'

    # Преобразуем файл doc в docx, т.к. библиотека не работает без этого
    w = wc.Dispatch('word.Application')
    doc_docx = w.Documents.Open(os.path.abspath(filepath))
    doc_docx.SaveAs(os.path.abspath(filepath) + "x", 16)
    doc_docx.Close()
    w.Quit()

    # filepath - финальный относительный путь до нужного документа
    filepath = f"data/{file_path[0]}" + 'x'



    # doc = docx.Document(filepath)
    #
    # all_paras = doc.paragraphs
    # # print(len(all_paras))
    #
    # for para in all_paras:
    #     print(para.text)
    #     print("-----")
    # # ///////////////////
    #
    # all_tables = doc.tables
    # # print(len(all_tables))
    #
    # data_tables = {i: None for i in range(len(all_tables))}
    # # проходимся по таблицам
    # for i, table in enumerate(all_tables):
    #     print('\nДанные таблицы №', i)
    #     # создаем список строк для таблицы `i` (пока пустые)
    #     data_tables[i] = [[] for _ in range(len(table.rows))]
    #     # проходимся по строкам таблицы `i`
    #     for j, row in enumerate(table.rows):
    #         # проходимся по ячейкам таблицы `i` и строки `j`
    #         for cell in row.cells:
    #             # добавляем значение ячейки в соответствующий
    #             # список, созданного словаря под данные таблиц
    #             data_tables[i][j].append(cell.text)
    #     print(data_tables[i])
    #     print('\n')
    #
    # print('Данные всех таблиц документа:')
    # print(data_tables)
    # # ///////////////////
    #
    # string_with_package = all_paras[0].text
    # package_format = packageformatdef(string_with_package)

    dict_push = docx_parser.docx_parse(filepath)

    if not os.path.isdir(dict_push["order"]):
        os.mkdir(dict_push["order"])

    if not os.path.isdir(dict_push["order"]+"/"+dict_push["block"]):
        os.mkdir(dict_push["order"]+"/"+dict_push["block"])

    if not os.path.isdir(dict_push["order"]+"/"+dict_push["block"]+"/"+dict_push["package"]):
        os.mkdir(dict_push["order"]+"/"+dict_push["block"]+"/"+dict_push["package"])

    curr_path = dict_push["order"]+"/"+dict_push["block"]+"/"+dict_push["package"]+"/AccDocs"

    if not os.path.isdir(curr_path):
        os.mkdir(curr_path)
    if "Check-list" in dict_push["typefile"]:
        if not os.path.isdir(curr_path+"/CheckList"):
            os.mkdir(curr_path+"/CheckList")
            createxml()
    elif "IKL" in dict_push["typefile"]:
        if not os.path.isdir(curr_path+"/IKL"):
            os.mkdir(curr_path+"/IKL")
            createxml()
    elif "Notes" or "Рабочая документация" in dict_push["typefile"]:
        if not os.path.isdir(curr_path+"/Notes"):
            os.mkdir(curr_path+"/Notes")
            createxml()
    elif "PDTK" in dict_push["typefile"]:
        if not os.path.isdir(curr_path+"/PDTK"):
            os.mkdir(curr_path+"/PDTK")
            createxml()
    # else:
    #     if not os.path.isdir(package_format + "/Docs"):
    #         os.mkdir(package_format + "/Docs")
    #         createxml()
    # print(datetime.datetime.fromtimestamp(os.path.getctime(filepath)).strftime('%Y%m%d%H%M%S'))
