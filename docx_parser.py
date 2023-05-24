import re
import docx


def find_type(name_type, text):
    typefile = re.search(name_type, text)
    typefile_cg = typefile[0] if typefile else 'Not found'
    return typefile_cg


def docx_parse(filepath):
    # Открытие и сбор текста из файла
    doc = docx.Document(filepath)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    full_text_table = []
    for table in doc.tables:
        for column in table.columns:
            for cell in column.cells:
                full_text_table.append(cell.text)
                # print(cell.text)
    text_paragraphs = ' '.join(text)
    text_tables = ' '.join(full_text_table)
    text_all = text_paragraphs + text_tables

    # Проверка типа документа
    typefile_cg = find_type("Рабочая документация", text_all)
    if typefile_cg == 'Not found':
        typefile_cg = find_type("Чек-лист", text_all)
    if typefile_cg == 'Not found':
        typefile_cg = find_type("Сопроводительное письмо", text_all)
    if typefile_cg == 'Not found':
        typefile_cg = find_type("Пояснительная записка", text_all)

    # Случай для главной ведомости
    if typefile_cg == "Рабочая документация":
        print("----------")

        # Создаём словарь для ведомости
        dict_info_main = {}

        dict_info_main["typefile"] = typefile_cg
        dict_info_main["document_id"] = "12345"
        # Собираем названия всех файлов из первой таблицы
        files_list = []
        for i in range(len(doc.tables[0].columns[0].cells)):
            if i != 0 and i != 1:
                files_list.append(doc.tables[0].columns[0].cells[i].text)

        print(files_list)
        dict_info_main["files_list"] = files_list
        # Собираем данные из второй таблицы
        for j in range(len(doc.tables[1].rows)):
            for i in range(len(doc.tables[1].rows[j].cells)):
                if j != 0:
                    print(doc.tables[1].rows[j].cells[i].text)
                    text_clear = doc.tables[1].rows[j].cells[i].text
                    text_clear = re.sub(r"[/\|\?]", '-', text_clear, count=0)
                    if i == 0:
                        dict_info_main["order"] = text_clear
                    elif i == 1:
                        dict_info_main["block"] = text_clear
                    elif i == 2:
                        dict_info_main["package"] = text_clear
                    else:
                        dict_info_main[doc.tables[1].rows[0].cells[i].text] = text_clear
                        # full_text_table.append(cell.text)
        print(dict_info_main)
        # Возвращаем словарь с информацией о ведомости
        return dict_info_main
    else:

        dict_info = {}
        dict_info["typefile"] = typefile_cg
        dict_info["list_other_column"] = []
        dict_info["id_element"] = []
        dict_info["id_work"] = []
        for table in doc.tables:
            for j in range(len(table.rows)):
                for i in range(len(table.rows[j].cells)):
                    if j != 0:
                        text_clear = table.rows[j].cells[i].text
                        text_clear = re.sub(r"[/\|\?]", '-', text_clear, count=0).rstrip().lstrip()
                        if i == 0:
                            dict_info["order"] = text_clear
                        elif i == 1:
                            dict_info["block"] = text_clear
                        elif i == 2:
                            dict_info["package"] = text_clear
                        elif i == 3:
                            dict_info["id_element"].append(text_clear)
                        elif i == 4:
                            dict_info["id_work"].append(text_clear)
                            dict_info["document_id"] = text_clear
                        else:
                            dict_info["list_other_column"].append(table.rows[0].cells[i].text)
                            dict_info[table.rows[0].cells[i].text].append(text_clear)
                    else:
                        if i > 4:
                            dict_info[table.rows[0].cells[i].text] = []

        print(dict_info)
        return dict_info


# if __name__ == '__main__':
#     docx_parse("data/RU_5_9_3_10.docx")
