import re
import docx

doc = docx.Document("Чек-лист _5 9 3 10 RUENG.docx")
print(doc.paragraphs)


text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
print('\n'.join(text))

full_text_table = []
for table in doc.tables:
    for column in table.columns:
        for cell in column.cells:
            full_text_table.append(cell.text)
            # print(cell.text)



test = ' '.join(full_text_table)

print(test)

order_match = re.search(r"\d{10}", test)
print(order_match[0] if order_match else 'Not found')

ID_match = re.search(r"\S{22}-\S{2,15}", test)
print(ID_match[0] if ID_match else 'Not found')

Designation_match = re.search(r"Designation \(link\)\s{1,}\S{2,} ", test)
print(Designation_match[0].replace('Designation (link) ','') if Designation_match else 'Not found')

Comment_match = re.search(r"Comment \(can be Null\)\s{1,}\S{2,}", test)
print(Comment_match[0].replace('Comment (can be Null) ','') if Comment_match else 'Not found')
